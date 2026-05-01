"""Generate laporan training dalam format .docx (Microsoft Word).

Usage:
    python scripts/generate_training_report.py
"""
from __future__ import annotations

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "2E4057")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if i % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return p


def add_note(doc, label: str, text: str):
    """Add a styled note/warning box (simulated with indented italic)."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"[{label}] ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x8B, 0x00, 0x00) if "Penting" in label or "Kritis" in label else RGBColor(0x00, 0x4E, 0x8C)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.italic = True
    p.paragraph_format.left_indent = Cm(1)
    return p


# ── Main document builder ───────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()

    # Page style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ═══════════════════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════════════════
    title = doc.add_heading("Laporan Proses Training Model Prediksi Muka Air Dhompo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Sistem Early Warning Banjir DAS Rejoso, Kabupaten Pasuruan")
    run.font.size = Pt(13)
    run.italic = True

    # ═══════════════════════════════════════════════════════════════════════
    # 1. PENDAHULUAN
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Pendahuluan")

    add_para(doc,
        "Dokumen ini menjelaskan secara lengkap proses pengembangan model prediksi muka air "
        "untuk stasiun Dhompo (elevasi 7 mdpl) pada DAS Rejoso, Kabupaten Pasuruan. "
        "Sistem ini dirancang sebagai komponen inti Early Warning System (EWS) banjir yang mampu "
        "memberikan prediksi 1 hingga 5 jam ke depan dengan resolusi temporal 30 menit."
    )

    add_heading(doc, "1.1 Tujuan", level=2)
    for t in [
        "Membangun model prediksi multi-horizon (+1h s.d. +5h) muka air di stasiun Dhompo.",
        "Mengidentifikasi algoritma dan konfigurasi fitur terbaik melalui serangkaian eksperimen sistematis.",
        "Mendiagnosis kelemahan model pada kondisi ekstrem (banjir) dan merumuskan strategi perbaikan.",
        "Merancang mekanisme penanganan data hilang (missing data) agar model tetap beroperasi saat sensor mengalami gangguan.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    add_heading(doc, "1.2 Pendekatan: Direct Multi-Horizon Forecasting", level=2)
    add_para(doc,
        "Sistem menggunakan strategi Direct Forecasting \u2014 setiap horizon prediksi (h1\u2013h5) "
        "memiliki model independen. Berbeda dengan pendekatan recursive yang memprediksi satu langkah "
        "lalu menggunakan hasil prediksi sebagai input berikutnya, direct forecasting menghindari "
        "akumulasi error antar horizon. Konsekuensinya, setiap horizon dapat menggunakan algoritma "
        "yang berbeda sesuai karakteristik prediksi masing-masing."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 2. DATA DAN PREPROCESSING
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Data dan Preprocessing")

    add_heading(doc, "2.1 Sumber Data", level=2)
    add_para(doc,
        "Sistem menggunakan dua sumber data yang digabungkan secara segment-aware:"
    )
    add_table(doc,
        ["Dataset", "Periode", "Durasi", "Catatan"],
        [
            ["2022 (Historis)", "Okt\u2013Des 2022", "~67 hari", "14 stasiun muka air, tanpa curah hujan"],
            ["2023 (Generated)", "Jan\u2013Mar 2023", "~90 hari", "14 stasiun muka air + curah hujan"],
        ],
    )

    add_heading(doc, "2.2 Jaringan Stasiun", level=2)
    add_para(doc,
        "Data berasal dari 14 stasiun pemantauan muka air yang tersebar sepanjang DAS dengan "
        "gradien elevasi dari hulu (503 mdpl) ke hilir (1.8 mdpl):"
    )
    add_table(doc,
        ["No", "Stasiun", "Elevasi (mdpl)", "Peran"],
        [
            ["1",  "Bd. Suwoto",       "503",  "Hulu"],
            ["2",  "Krajan Timur",     "335",  "Hulu"],
            ["3",  "Purwodadi",        "287",  "Hulu"],
            ["4",  "Bd. Baong",        "169",  "Tengah"],
            ["5",  "Bd. Lecari",       "167",  "Tengah"],
            ["6",  "Bd. Bakalan",      "136",  "Tengah"],
            ["7",  "AWLR Kademungan",  "128",  "Tengah"],
            ["8",  "Bd. Domas",        "57",   "Hilir"],
            ["9",  "Bd Guyangan",      "32",   "Hilir"],
            ["10", "Bd. Grinting",     "28",   "Hilir"],
            ["11", "Sidogiri",         "24",   "Hilir"],
            ["12", "Klosod",           "22",   "Hilir"],
            ["13", "Dhompo (Target)",  "7",    "Target"],
            ["14", "Jalan Nasional",   "1.8",  "Muara"],
        ],
    )
    add_para(doc,
        "Model menerima input dari 12 stasiun hulu (No. 1\u201312) ditambah Dhompo sendiri "
        "sebagai prediktor, dan memprediksi muka air di Dhompo pada waktu mendatang."
    )

    add_heading(doc, "2.3 Validasi Kompatibilitas Dataset", level=2)
    add_para(doc,
        "Sebelum penggabungan, dilakukan validasi kompatibilitas antara dataset 2022 dan 2023:"
    )
    for t in [
        "Kolmogorov-Smirnov test mengkonfirmasi perbedaan musiman yang diharapkan (musim kemarau vs musim hujan).",
        "Profil statistik (mean, std, range) konsisten secara fisik \u2014 dataset 2023 memiliki variabilitas lebih tinggi karena musim hujan.",
        "Curah hujan hanya tersedia pada dataset 2023; untuk dataset 2022 diasumsikan kering (diisi 0).",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    add_heading(doc, "2.4 Strategi Split Data", level=2)
    add_para(doc,
        "Seluruh data 2022 digunakan untuk training. Data 2023 dibagi 80/20 secara temporal "
        "(20% terakhir sebagai test set). Temporal split tanpa shuffling menjaga kronologi waktu "
        "dan mencegah data leakage."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 3. REKAYASA FITUR
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Rekayasa Fitur")

    add_heading(doc, "3.1 Fitur Baseline (160 fitur)", level=2)
    add_para(doc,
        "Setiap timestep direpresentasikan oleh fitur yang dibangun dari 13 stasiun "
        "(12 hulu + Dhompo) ditambah fitur temporal:"
    )
    add_table(doc,
        ["Kategori", "Deskripsi", "Jumlah"],
        [
            ["Nilai saat ini (t0)", "Muka air pada waktu sekarang", "13"],
            ["Lag (t-1, t-2, t-3)", "Nilai 30, 60, 90 menit yang lalu", "39"],
            ["Rolling mean", "Rata-rata bergerak 3h, 6h, 12h per stasiun", "39"],
            ["Rolling std", "Standar deviasi bergerak 3h, 6h, 12h per stasiun", "39"],
            ["Rate of change", "diff1 (\u039430 menit), diff2 (\u039460 menit) per stasiun", "26"],
            ["Temporal", "hour_sin, hour_cos, dayofweek, is_night", "4"],
            ["Total", "", "160"],
        ],
    )

    add_heading(doc, "3.2 Fitur Tambahan (Progressive Features)", level=2)
    add_para(doc, "Fitur tambahan diuji secara inkremental di atas baseline:")
    add_table(doc,
        ["Kode", "Fitur Tambahan", "Deskripsi"],
        [
            ["A",  "Baseline + Rainfall", "Baseline 160 fitur + fitur curah hujan (t0, lag, rolling, diff)"],
            ["B1", "+ Travel Time Lags", "Lag berbasis travel time empiris per stasiun (bukan lag seragam 1-3)"],
            ["B2", "+ Cumulative Rainfall", "Rolling sum curah hujan 3h/6h/12h/24h"],
            ["B3", "+ Interaction", "Gradien hidraulik antar stasiun + interaksi gradien\u00d7curah hujan"],
            ["B4", "Full Features", "Semua fitur di atas + musim basah + antecedent moisture index 7 hari"],
        ],
    )

    add_heading(doc, "3.3 Preprocessing", level=2)
    for t in [
        "StandardScaler diterapkan hanya untuk model linear (Ridge, Lasso, ElasticNet) yang sensitif terhadap skala.",
        "Model tree-based (Random Forest, Gradient Boosting, XGBoost, LightGBM, CatBoost) menggunakan fitur tanpa scaling.",
        "Minimum data: dibutuhkan 24 baris (12 jam) data historis agar semua fitur rolling terisi.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ═══════════════════════════════════════════════════════════════════════
    # 4. EKSPERIMEN DAN HASIL
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Eksperimen dan Hasil")

    add_para(doc,
        "Proses pengembangan model dilakukan melalui serangkaian eksperimen sistematis. "
        "Setiap eksperimen menguji satu aspek tertentu (data, fitur, target, atau strategi training) "
        "sambil menjaga variabel lain tetap konstan. Berikut ringkasan seluruh eksperimen yang dilakukan:"
    )

    add_table(doc,
        ["No", "Eksperimen", "Aspek yang Diuji", "Deskripsi"],
        [
            ["1", "Training 2022 Only",         "Data",     "Model dilatih hanya pada data 2022 (Okt\u2013Des), dievaluasi pada data 2022 itu sendiri"],
            ["2", "A: Train 2022 \u2192 Test 2023", "Generalisasi", "Model dari data 2022 diuji langsung pada data 2023 tanpa fine-tuning"],
            ["3", "B: Combined Training",        "Data",     "Gabung data 2022 + 2023, temporal split 80/20 pada 2023"],
            ["4", "C: Combined + Rainfall",      "Fitur",    "Sama dengan B, ditambah fitur curah hujan dari dataset 2023"],
            ["5", "Progressive Features (A\u2013B4)", "Fitur", "Fitur tambahan secara inkremental: travel time lags, cumulative rain, interaction, seasonal"],
            ["6", "Delta vs Absolute Target",    "Target",   "Membandingkan prediksi perubahan (\u0394y) vs nilai absolut (y)"],
            ["7", "Target Smoothing",            "Target",   "Meng-smooth target dengan moving average/median sebelum training"],
        ],
    )

    # 4.1 Model yang Dievaluasi
    add_heading(doc, "4.1 Algoritma Model yang Dievaluasi", level=2)
    add_para(doc,
        "Setiap eksperimen mengevaluasi hingga 9 algoritma machine learning secara paralel:"
    )
    add_table(doc,
        ["No", "Algoritma", "Tipe", "Scaling"],
        [
            ["1", "Linear Regression", "Linear",         "Ya"],
            ["2", "Ridge",             "Linear (L2)",    "Ya"],
            ["3", "Lasso",             "Linear (L1)",    "Ya"],
            ["4", "ElasticNet",        "Linear (L1+L2)", "Ya"],
            ["5", "Random Forest",     "Ensemble",       "Tidak"],
            ["6", "Gradient Boosting", "Ensemble",       "Tidak"],
            ["7", "XGBoost",           "Ensemble",       "Tidak"],
            ["8", "LightGBM",          "Ensemble",       "Tidak"],
            ["9", "CatBoost",          "Ensemble",       "Tidak"],
        ],
    )
    add_para(doc,
        "Model tree-based (No. 5\u20139) tidak memerlukan scaling fitur, sedangkan model linear "
        "(No. 1\u20134) menggunakan StandardScaler. Berikut penjelasan arsitektur masing-masing algoritma."
    )

    # Deskripsi arsitektur tiap model
    add_heading(doc, "A. Linear Regression", level=3)
    add_para(doc,
        "Model regresi paling dasar yang mencari hubungan linear antara fitur input dan target "
        "dengan meminimalkan jumlah kuadrat residual (Ordinary Least Squares). "
        "Setiap fitur diberi koefisien bobot, dan prediksi dihasilkan dari penjumlahan "
        "linear seluruh fitur berbobot ditambah intercept. Model ini digunakan sebagai baseline "
        "karena kesederhanaannya, namun tidak mampu menangkap hubungan non-linear antar variabel."
    )

    add_heading(doc, "B. Ridge Regression", level=3)
    add_para(doc,
        "Pengembangan dari Linear Regression dengan penambahan regularisasi L2 pada fungsi loss. "
        "Regularisasi L2 menambahkan penalti sebesar \u03b1\u2211\u03b2\u00b2 (jumlah kuadrat koefisien) "
        "yang mencegah koefisien tumbuh terlalu besar. Hal ini mengurangi overfitting, terutama "
        "saat jumlah fitur besar atau terdapat multikolinearitas antar fitur. Berbeda dengan Lasso, "
        "Ridge tidak mengeliminasi fitur \u2014 semua fitur tetap digunakan namun dengan bobot yang dikecilkan."
    )

    add_heading(doc, "C. Lasso Regression", level=3)
    add_para(doc,
        "Regresi linear dengan regularisasi L1 yang menambahkan penalti sebesar \u03b1\u2211|\u03b2| "
        "(jumlah nilai absolut koefisien). Keunikan Lasso adalah kemampuannya mendorong koefisien "
        "fitur yang kurang penting menjadi tepat nol, sehingga secara otomatis melakukan seleksi fitur. "
        "Hal ini berguna saat banyak fitur yang tidak relevan, karena model yang dihasilkan lebih "
        "sparse dan mudah diinterpretasi."
    )

    add_heading(doc, "D. ElasticNet", level=3)
    add_para(doc,
        "Kombinasi dari regularisasi L1 (Lasso) dan L2 (Ridge) dalam satu fungsi loss. "
        "ElasticNet mengatasi keterbatasan Lasso yang cenderung memilih hanya satu fitur secara "
        "arbitrer dari sekelompok fitur yang berkorelasi tinggi. Dengan menggabungkan kedua penalti, "
        "ElasticNet dapat melakukan seleksi fitur (sifat L1) sekaligus menangani multikolinearitas "
        "(sifat L2). Rasio antara L1 dan L2 diatur oleh parameter l1_ratio."
    )

    add_heading(doc, "E. Random Forest", level=3)
    add_para(doc,
        "Algoritma ensemble berbasis bagging (Bootstrap Aggregating) yang membangun banyak decision tree "
        "secara paralel dan independen. Setiap tree dilatih pada subset acak dari data training "
        "(bootstrap sampling) dan subset acak dari fitur (feature bagging). Prediksi final adalah "
        "rata-rata dari prediksi seluruh tree. Randomisasi ganda ini menghasilkan model yang robust "
        "terhadap outlier dan noise, serta mengurangi variansi prediksi dibanding single decision tree. "
        "Tidak memerlukan scaling fitur karena decision tree bersifat invariant terhadap skala."
    )

    add_heading(doc, "F. Gradient Boosting", level=3)
    add_para(doc,
        "Algoritma ensemble berbasis boosting yang membangun decision tree secara sekuensial. "
        "Berbeda dengan Random Forest yang membangun tree secara paralel, Gradient Boosting "
        "membangun setiap tree baru untuk memperbaiki residual (error) dari tree-tree sebelumnya. "
        "Proses ini mengikuti arah gradien negatif dari fungsi loss, sehingga setiap iterasi "
        "secara bertahap mengurangi error. Learning rate mengontrol kontribusi setiap tree baru "
        "untuk mencegah overfitting. Model ini unggul dalam menangkap pola non-linear yang kompleks."
    )

    add_heading(doc, "G. XGBoost (Extreme Gradient Boosting)", level=3)
    add_para(doc,
        "Implementasi gradient boosting yang dioptimasi dengan beberapa peningkatan arsitektural: "
        "(1) regularisasi L1 dan L2 bawaan pada bobot leaf node untuk mencegah overfitting, "
        "(2) approximate tree learning menggunakan histogram-based split finding yang mempercepat training, "
        "(3) penanganan missing values secara native dengan mempelajari arah split optimal, "
        "dan (4) column subsampling per tree dan per level yang mirip Random Forest. "
        "Kombinasi ini menghasilkan model yang akurat, cepat, dan robust."
    )

    add_heading(doc, "H. LightGBM (Light Gradient Boosting Machine)", level=3)
    add_para(doc,
        "Varian gradient boosting yang menggunakan dua teknik utama untuk efisiensi: "
        "(1) Gradient-based One-Side Sampling (GOSS), yang mempertahankan sampel dengan gradien besar "
        "dan hanya mengambil sebagian sampel bergradien kecil, sehingga fokus pada data yang sulit diprediksi; "
        "(2) Exclusive Feature Bundling (EFB), yang menggabungkan fitur-fitur yang jarang aktif bersamaan "
        "untuk mengurangi dimensi. LightGBM juga menggunakan strategi pertumbuhan leaf-wise "
        "(memilih leaf dengan loss reduction terbesar) alih-alih level-wise, menghasilkan tree yang "
        "lebih dalam namun lebih akurat dengan jumlah leaf yang sama."
    )

    add_heading(doc, "I. CatBoost (Categorical Boosting)", level=3)
    add_para(doc,
        "Algoritma gradient boosting yang dirancang untuk menangani fitur kategorikal secara native "
        "dan mengurangi target leakage melalui teknik ordered boosting. Pada ordered boosting, "
        "setiap sampel training hanya melihat residual dari sampel-sampel yang diproses sebelumnya "
        "(berdasarkan permutasi acak), mencegah model mempelajari target secara langsung. "
        "CatBoost juga menggunakan oblivious decision tree (symmetric tree) di mana setiap level "
        "menggunakan split condition yang sama untuk semua node, menghasilkan model yang lebih cepat "
        "saat inferensi dan lebih tahan terhadap overfitting. Dalam proyek ini, CatBoost secara konsisten "
        "menjadi algoritma dengan performa terbaik di hampir semua horizon dan konfigurasi eksperimen."
    )

    # 4.2 Training 2022 Only
    add_heading(doc, "4.2 Eksperimen 1: Training Data 2022 Only", level=2)

    add_para(doc, "Tujuan:", bold=True)
    add_para(doc,
        "Membangun model prediksi awal menggunakan data historis yang tersedia "
        "dan mengevaluasi seberapa baik berbagai algoritma mampu mempelajari pola muka air."
    )

    add_para(doc, "Flow eksperimen:", bold=True)
    for t in [
        "Data yang digunakan: data historis Oktober\u2013Desember 2022 (3.135 baris, 14 stasiun, tanpa curah hujan).",
        "Rekayasa fitur: 160 fitur baseline (nilai saat ini, lag 1\u20133, rolling mean/std 3h/6h/12h, diff1/diff2, fitur temporal).",
        "Split data: temporal split 80% training / 20% test (tanpa shuffle, menjaga urutan waktu).",
        "Training: 6 algoritma (Linear Regression, Ridge, Lasso, Random Forest, Gradient Boosting, XGBoost) dilatih secara independen untuk masing-masing 5 horizon.",
        "Evaluasi: metrik NSE, RMSE, dan MAE dihitung pada test set.",
    ]:
        doc.add_paragraph(t, style="List Number")

    add_para(doc, "Hasil NSE dan RMSE per model per horizon:", bold=True)
    # 2022-only: per horizon table
    for h_label, h_rows in [
        ("Horizon +1 Jam", [
            ["Linear Regression", "0.9819", "0.1642"],
            ["Ridge",             "0.9827", "0.1607"],
            ["Lasso",             "0.9807", "0.1698"],
            ["Random Forest",     "0.9875", "0.1368"],
            ["Gradient Boosting", "0.9891", "0.1275"],
            ["XGBoost",           "0.9897", "0.1242"],
        ]),
        ("Horizon +2 Jam", [
            ["Linear Regression", "0.9771", "0.1849"],
            ["Ridge",             "0.9776", "0.1831"],
            ["Lasso",             "0.9742", "0.1967"],
            ["Random Forest",     "0.9803", "0.1715"],
            ["Gradient Boosting", "0.9825", "0.1618"],
            ["XGBoost",           "0.9825", "0.1620"],
        ]),
        ("Horizon +3 Jam", [
            ["Linear Regression", "0.9500", "0.2739"],
            ["Ridge",             "0.9525", "0.2670"],
            ["Lasso",             "0.9525", "0.2669"],
            ["Random Forest",     "0.9525", "0.2671"],
            ["Gradient Boosting", "0.9563", "0.2562"],
            ["XGBoost",           "0.9542", "0.2621"],
        ]),
        ("Horizon +4 Jam", [
            ["Linear Regression", "0.8832", "0.4193"],
            ["Ridge",             "0.8849", "0.4163"],
            ["Lasso",             "0.8894", "0.4081"],
            ["Random Forest",     "0.8835", "0.4189"],
            ["Gradient Boosting", "0.8806", "0.4241"],
            ["XGBoost",           "0.8815", "0.4224"],
        ]),
        ("Horizon +5 Jam", [
            ["Linear Regression", "0.7539", "0.6097"],
            ["Ridge",             "0.7590", "0.6034"],
            ["Lasso",             "0.7713", "0.5877"],
            ["Random Forest",     "0.7650", "0.5959"],
            ["Gradient Boosting", "0.7563", "0.6068"],
            ["XGBoost",           "0.7644", "0.5965"],
        ]),
    ]:
        add_para(doc, h_label, bold=True, font_size=10)
        add_table(doc, ["Model", "NSE", "RMSE (m)"], h_rows)

    add_note(doc, "Catatan",
        "Performa pada data 2022 sendiri sangat baik (NSE > 0.98 pada h1\u2013h2) karena test set "
        "berasal dari periode dan distribusi yang sama dengan training set. Namun hal ini belum "
        "menunjukkan kemampuan generalisasi ke periode lain."
    )

    # 4.3 A vs B vs C
    add_heading(doc, "4.3 Eksperimen 2\u20134: Strategi Penggabungan Data (A vs B vs C)", level=2)

    add_para(doc, "Tujuan:", bold=True)
    add_para(doc,
        "Menentukan strategi terbaik untuk memanfaatkan dua dataset yang tersedia (2022 dan 2023) "
        "dan menguji apakah penambahan fitur curah hujan meningkatkan akurasi."
    )

    add_para(doc, "Flow eksperimen:", bold=True)
    for t in [
        "Eksperimen A \u2014 Uji generalisasi: Model yang sudah dilatih pada data 2022 (dari Eksperimen 1) "
        "langsung diuji pada data 2023 tanpa training ulang. Tujuannya mengukur seberapa mampu model "
        "yang dilatih di musim kemarau menangani dinamika musim hujan.",
        "Eksperimen B \u2014 Combined training: Data 2022 dan 2023 digabungkan secara segment-aware "
        "(mencegah kontaminasi fitur rolling antar gap waktu). Seluruh data 2022 masuk training, "
        "data 2023 dibagi 80% training / 20% test secara temporal. "
        "Semua 9 algoritma dilatih dan dievaluasi ulang.",
        "Eksperimen C \u2014 Combined + Rainfall: Sama dengan B, namun fitur curah hujan dari dataset 2023 "
        "ditambahkan ke feature matrix (t0, lag1\u20133, rolling mean/std 3h/6h/12h, diff1/diff2). "
        "Untuk data 2022 yang tidak memiliki curah hujan, kolom diisi 0 (asumsi kering).",
        "Evaluasi: Ketiga eksperimen dibandingkan menggunakan model terbaik per horizon.",
    ]:
        doc.add_paragraph(t, style="List Number")

    add_table(doc,
        ["Eksperimen", "Strategi", "Deskripsi"],
        [
            ["A", "Train 2022 \u2192 Test 2023", "Uji generalisasi langsung tanpa fine-tuning"],
            ["B", "Combined training", "Gabung data 2022 + 2023, temporal split pada 2023"],
            ["C", "Combined + Rainfall", "Sama dengan B, ditambah fitur curah hujan"],
        ],
    )

    add_para(doc, "Hasil perbandingan (model terbaik per horizon):", bold=True)
    add_table(doc,
        ["Horizon", "A: NSE", "A: RMSE", "A: Model", "B: NSE", "B: RMSE", "B: Model", "C: NSE", "C: RMSE", "C: Model"],
        [
            ["+1h", "0.765",    "0.330", "Grad. Boosting", "0.854", "0.180", "Random Forest",     "0.854", "0.180", "Random Forest"],
            ["+2h", "0.269",    "0.579", "Grad. Boosting", "0.840", "0.189", "Ridge",             "0.840", "0.189", "Grad. Boosting"],
            ["+3h", "\u22125.093", "1.668", "Grad. Boosting", "0.819", "0.201", "Ridge",           "0.818", "0.201", "Ridge"],
            ["+4h", "\u221221.526","3.199", "Grad. Boosting", "0.795", "0.214", "Ridge",           "0.793", "0.215", "Ridge"],
            ["+5h", "\u221220.218","3.096", "XGBoost",        "0.751", "0.236", "Lasso",           "0.751", "0.236", "Lasso"],
        ],
    )

    add_note(doc, "Temuan Kritis",
        "Eksperimen A menunjukkan kegagalan total pada horizon h3\u2013h5 (NSE negatif, artinya model "
        "lebih buruk dari rata-rata). Model yang dilatih hanya pada data musim kemarau 2022 tidak mampu "
        "menangani dinamika musim hujan 2023. Combined training (B) menyelesaikan masalah ini secara dramatis."
    )
    add_note(doc, "Catatan",
        "Curah hujan (Eksperimen C) hanya memberikan perbaikan marginal dibanding B. Sinyal curah hujan "
        "sudah ter-capture secara tidak langsung melalui kenaikan muka air di stasiun hulu."
    )

    # 4.4 Progressive Features — all models per experiment
    add_heading(doc, "4.4 Eksperimen 5: Progressive Features", level=2)

    add_para(doc, "Tujuan:", bold=True)
    add_para(doc,
        "Menguji apakah penambahan fitur yang lebih canggih (berdasarkan domain knowledge hidrologi) "
        "dapat meningkatkan akurasi prediksi di atas baseline."
    )

    add_para(doc, "Flow eksperimen:", bold=True)
    for t in [
        "Baseline ditetapkan: konfigurasi A (160 fitur + rainfall) dari Eksperimen C yang sudah terbukti "
        "menghasilkan combined training terbaik.",
        "Fitur ditambahkan secara inkremental: B1 menambahkan travel time lags (lag berbasis waktu tempuh "
        "empiris per stasiun, bukan lag seragam 1\u20133); B2 menambahkan cumulative rainfall "
        "(rolling sum curah hujan 3h/6h/12h/24h); B3 menambahkan interaction features "
        "(gradien hidraulik antar stasiun berurutan + interaksi gradien\u00d7curah hujan); "
        "B4 menggabungkan semua fitur di atas plus seasonal features (flag musim basah + antecedent moisture index 7 hari).",
        "Setiap konfigurasi fitur (A, B1, B2, B3, B4) dilatih menggunakan seluruh 9 algoritma "
        "pada data combined (2022+2023) dengan split yang sama.",
        "Evaluasi: NSE dan RMSE dihitung pada test set untuk setiap kombinasi konfigurasi fitur \u00d7 algoritma \u00d7 horizon.",
    ]:
        doc.add_paragraph(t, style="List Number")

    add_para(doc,
        "Berikut hasil NSE dan RMSE setiap model untuk masing-masing konfigurasi fitur, per horizon."
    )

    # -- Data: Progressive Features all models --
    _pf_models = [
        "Linear Regression", "Ridge", "Lasso", "ElasticNet",
        "Random Forest", "Gradient Boosting", "XGBoost", "LightGBM", "CatBoost",
    ]
    _pf_data = {
        "A (Baseline + Rainfall)": {
            1: [("0.849","0.183"),("0.850","0.183"),("0.837","0.191"),("0.841","0.188"),("0.854","0.180"),("0.850","0.183"),("0.852","0.182"),("0.859","0.177"),("0.861","0.176")],
            2: [("0.838","0.190"),("0.839","0.189"),("0.830","0.194"),("0.837","0.191"),("0.834","0.192"),("0.840","0.189"),("0.836","0.191"),("0.841","0.188"),("0.848","0.184")],
            3: [("0.817","0.202"),("0.818","0.201"),("0.804","0.209"),("0.811","0.205"),("0.801","0.210"),("0.807","0.207"),("0.814","0.203"),("0.810","0.206"),("0.834","0.192")],
            4: [("0.792","0.215"),("0.793","0.215"),("0.783","0.220"),("0.788","0.218"),("0.714","0.253"),("0.733","0.244"),("0.747","0.238"),("0.773","0.225"),("0.793","0.215")],
            5: [("0.745","0.239"),("0.743","0.240"),("0.751","0.236"),("0.754","0.235"),("0.651","0.279"),("0.637","0.285"),("0.674","0.270"),("0.680","0.268"),("0.749","0.237")],
        },
        "B1 (+ Travel Time Lags)": {
            1: [("0.847","0.185"),("0.848","0.184"),("0.836","0.191"),("0.840","0.189"),("0.853","0.181"),("0.844","0.186"),("0.847","0.185"),("0.858","0.178"),("0.857","0.179")],
            2: [("0.840","0.189"),("0.841","0.188"),("0.830","0.194"),("0.836","0.191"),("0.831","0.194"),("0.830","0.194"),("0.828","0.196"),("0.835","0.192"),("0.845","0.186")],
            3: [("0.817","0.202"),("0.819","0.201"),("0.804","0.208"),("0.812","0.205"),("0.798","0.212"),("0.807","0.207"),("0.812","0.204"),("0.817","0.202"),("0.831","0.194")],
            4: [("0.790","0.216"),("0.792","0.215"),("0.783","0.220"),("0.788","0.217"),("0.716","0.252"),("0.743","0.240"),("0.753","0.235"),("0.764","0.229"),("0.802","0.210")],
            5: [("0.741","0.240"),("0.740","0.241"),("0.752","0.235"),("0.754","0.234"),("0.637","0.285"),("0.632","0.287"),("0.694","0.262"),("0.676","0.269"),("0.748","0.237")],
        },
        "B2 (+ Cumulative Rainfall)": {
            1: [("0.847","0.185"),("0.848","0.184"),("0.836","0.191"),("0.840","0.189"),("0.854","0.180"),("0.847","0.185"),("0.852","0.182"),("0.856","0.179"),("0.859","0.177")],
            2: [("0.840","0.189"),("0.841","0.188"),("0.830","0.194"),("0.836","0.191"),("0.829","0.195"),("0.827","0.196"),("0.827","0.196"),("0.839","0.189"),("0.848","0.184")],
            3: [("0.817","0.202"),("0.819","0.201"),("0.804","0.208"),("0.812","0.205"),("0.801","0.210"),("0.803","0.209"),("0.817","0.202"),("0.814","0.204"),("0.833","0.193")],
            4: [("0.791","0.216"),("0.793","0.215"),("0.783","0.220"),("0.788","0.217"),("0.717","0.251"),("0.750","0.236"),("0.742","0.240"),("0.772","0.225"),("0.791","0.216")],
            5: [("0.744","0.239"),("0.743","0.240"),("0.752","0.235"),("0.754","0.234"),("0.650","0.280"),("0.651","0.279"),("0.690","0.263"),("0.697","0.260"),("0.744","0.239")],
        },
        "B3 (+ Interaction)": {
            1: [("0.847","0.185"),("0.848","0.184"),("0.836","0.191"),("0.841","0.188"),("0.854","0.181"),("0.844","0.187"),("0.850","0.183"),("0.860","0.177"),("0.863","0.175")],
            2: [("0.840","0.189"),("0.841","0.188"),("0.830","0.194"),("0.836","0.191"),("0.830","0.195"),("0.819","0.201"),("0.833","0.193"),("0.838","0.190"),("0.850","0.183")],
            3: [("0.817","0.202"),("0.819","0.201"),("0.805","0.208"),("0.812","0.204"),("0.798","0.212"),("0.811","0.205"),("0.809","0.206"),("0.812","0.205"),("0.831","0.194")],
            4: [("0.791","0.216"),("0.793","0.215"),("0.784","0.219"),("0.789","0.217"),("0.711","0.254"),("0.740","0.241"),("0.756","0.233"),("0.763","0.230"),("0.798","0.212")],
            5: [("0.744","0.239"),("0.743","0.240"),("0.753","0.235"),("0.754","0.234"),("0.645","0.281"),("0.638","0.285"),("0.717","0.251"),("0.696","0.260"),("0.727","0.247")],
        },
        "B4 (Full Features)": {
            1: [("0.847","0.184"),("0.848","0.184"),("0.836","0.191"),("0.841","0.188"),("0.855","0.180"),("0.840","0.189"),("0.853","0.181"),("0.858","0.178"),("0.862","0.176")],
            2: [("0.840","0.189"),("0.841","0.188"),("0.830","0.194"),("0.836","0.191"),("0.830","0.195"),("0.826","0.197"),("0.835","0.192"),("0.835","0.192"),("0.850","0.182")],
            3: [("0.817","0.202"),("0.819","0.201"),("0.805","0.208"),("0.812","0.204"),("0.799","0.211"),("0.809","0.206"),("0.810","0.206"),("0.812","0.204"),("0.830","0.195")],
            4: [("0.791","0.216"),("0.794","0.215"),("0.784","0.219"),("0.789","0.217"),("0.710","0.254"),("0.761","0.231"),("0.755","0.234"),("0.772","0.225"),("0.796","0.213")],
            5: [("0.744","0.239"),("0.744","0.239"),("0.753","0.235"),("0.754","0.234"),("0.641","0.283"),("0.635","0.286"),("0.700","0.259"),("0.706","0.256"),("0.735","0.243")],
        },
    }

    for exp_name, exp_horizons in _pf_data.items():
        add_para(doc, exp_name, bold=True, font_size=10)
        for h in [1, 2, 3, 4, 5]:
            vals = exp_horizons[h]
            rows = []
            for m_idx, model_name in enumerate(_pf_models):
                nse, rmse = vals[m_idx]
                rows.append([model_name, nse, rmse])
            add_para(doc, f"  Horizon +{h} Jam", bold=True, font_size=10)
            add_table(doc, ["Model", "NSE", "RMSE (m)"], rows)

    add_note(doc, "Catatan",
        "Fitur tambahan memberikan perbaikan kecil di horizon pendek (h1\u2013h2) tetapi tidak konsisten "
        "di horizon panjang (h3\u2013h5). Baseline A sudah menangkap sebagian besar informasi prediktif. "
        "CatBoost secara konsisten menjadi model terbaik atau setara terbaik di hampir semua horizon dan konfigurasi fitur."
    )

    # 4.5 Delta vs Absolute — all models
    add_heading(doc, "4.5 Eksperimen 6: Delta vs Absolute Target", level=2)

    add_para(doc, "Tujuan:", bold=True)
    add_para(doc,
        "Menguji apakah mengubah representasi target dapat meningkatkan akurasi. "
        "Alih-alih memprediksi nilai absolut muka air y(t+h), model dilatih untuk memprediksi "
        "perubahan (delta) muka air \u0394y = y(t+h) \u2212 y(t)."
    )

    add_para(doc, "Flow eksperimen:", bold=True)
    for t in [
        "Mode ABS (Absolute): Target adalah nilai muka air di waktu mendatang y(t+h). "
        "Ini merupakan pendekatan standar yang digunakan pada eksperimen sebelumnya.",
        "Mode DELTA: Target diubah menjadi selisih \u0394y = y(t+h) \u2212 y(t). "
        "Saat inferensi, prediksi dikembalikan ke skala absolut dengan menambahkan nilai saat ini: "
        "y\u0302(t+h) = y(t) + \u0394y\u0302.",
        "Kedua mode dilatih pada data combined (2022+2023) dengan split dan fitur yang sama (Baseline A). "
        "Seluruh 9 algoritma dievaluasi.",
        "Evaluasi: NSE dan RMSE dihitung terhadap nilai absolut aktual (bukan terhadap delta), "
        "sehingga kedua mode dapat dibandingkan secara fair.",
    ]:
        doc.add_paragraph(t, style="List Number")

    add_para(doc, "Hasil per model per horizon:", bold=True)
    _da_models = [
        "Linear Regression", "Ridge", "Lasso", "ElasticNet",
        "Random Forest", "Gradient Boosting", "XGBoost", "LightGBM", "CatBoost",
    ]
    _da_data = {
        1: [("0.850","0.183","0.850","0.183"),("0.850","0.183","0.850","0.183"),("0.837","0.191","0.817","0.202"),("0.841","0.188","0.824","0.198"),("0.853","0.181","0.840","0.189"),("0.853","0.181","0.849","0.183"),("0.846","0.186","0.853","0.181"),("0.855","0.180","0.843","0.187"),("0.864","0.174","0.843","0.187")],
        2: [("0.839","0.190","0.839","0.190"),("0.840","0.189","0.839","0.189"),("0.830","0.194","0.812","0.205"),("0.837","0.191","0.819","0.201"),("0.833","0.193","0.822","0.199"),("0.834","0.192","0.815","0.203"),("0.825","0.197","0.822","0.199"),("0.842","0.187","0.828","0.196"),("0.848","0.184","0.821","0.199")],
        3: [("0.818","0.201","0.818","0.201"),("0.819","0.201","0.819","0.201"),("0.804","0.209","0.789","0.217"),("0.811","0.205","0.797","0.212"),("0.802","0.210","0.804","0.209"),("0.816","0.202","0.791","0.215"),("0.807","0.207","0.797","0.213"),("0.811","0.205","0.801","0.211"),("0.832","0.193","0.813","0.204")],
        4: [("0.794","0.214","0.794","0.214"),("0.795","0.214","0.795","0.214"),("0.783","0.220","0.766","0.229"),("0.788","0.217","0.774","0.225"),("0.714","0.253","0.751","0.235"),("0.732","0.245","0.753","0.235"),("0.739","0.241","0.707","0.256"),("0.759","0.232","0.772","0.226"),("0.799","0.212","0.779","0.222")],
        5: [("0.747","0.238","0.747","0.238"),("0.744","0.239","0.743","0.239"),("0.751","0.236","0.728","0.247"),("0.754","0.234","0.734","0.244"),("0.645","0.282","0.675","0.270"),("0.618","0.292","0.634","0.286"),("0.657","0.277","0.664","0.274"),("0.673","0.270","0.678","0.268"),("0.740","0.241","0.691","0.263")],
    }
    for h in [1, 2, 3, 4, 5]:
        add_para(doc, f"Horizon +{h} Jam", bold=True, font_size=10)
        rows = []
        for m_idx, model_name in enumerate(_da_models):
            abs_nse, abs_rmse, del_nse, del_rmse = _da_data[h][m_idx]
            rows.append([model_name, abs_nse, abs_rmse, del_nse, del_rmse])
        add_table(doc, ["Model", "ABS NSE", "ABS RMSE", "DELTA NSE", "DELTA RMSE"], rows)

    add_para(doc,
        "Kesimpulan: Prediksi nilai absolut konsisten lebih baik di semua horizon untuk hampir semua model. "
        "Delta target justru menurunkan performa, kemungkinan karena distribusi delta yang "
        "mendekati nol membuat model kesulitan membedakan perubahan kecil yang bermakna."
    )

    # 4.6 Target Smoothing — all models
    add_heading(doc, "4.6 Eksperimen 7: Target Smoothing", level=2)

    add_para(doc, "Tujuan:", bold=True)
    add_para(doc,
        "Menguji apakah meng-smooth target sebelum training dapat mengurangi noise pada data sensor "
        "dan membantu model mempelajari tren yang lebih halus, terutama di horizon panjang."
    )

    add_para(doc, "Flow eksperimen:", bold=True)
    for t in [
        "Empat konfigurasi target diuji: RAW (tanpa smoothing, sebagai baseline), "
        "MEAN3 (moving average window 3 timestep = 90 menit), "
        "MED3 (moving median window 3 timestep), dan MED5 (moving median window 5 timestep = 150 menit).",
        "Smoothing diterapkan hanya pada target (y) saat training. Fitur input (X) tidak di-smooth.",
        "Evaluasi dilakukan terhadap nilai raw aktual (bukan terhadap nilai yang di-smooth), "
        "sehingga semua konfigurasi dibandingkan pada ground truth yang sama.",
        "Seluruh 9 algoritma dilatih pada data combined (2022+2023) dengan split dan fitur Baseline A yang sama.",
    ]:
        doc.add_paragraph(t, style="List Number")

    add_para(doc, "Hasil per model per horizon:", bold=True)
    _sm_models = _da_models  # same 9 models
    _sm_data = {
        1: [
            ("0.849","0.183","0.843","0.187","0.844","0.186","0.832","0.194"),
            ("0.850","0.183","0.837","0.190","0.837","0.190","0.833","0.193"),
            ("0.837","0.191","0.827","0.197","0.822","0.199","0.819","0.201"),
            ("0.841","0.188","0.830","0.194","0.827","0.197","0.825","0.198"),
            ("0.854","0.180","0.847","0.185","0.845","0.186","0.846","0.185"),
            ("0.850","0.183","0.850","0.183","0.846","0.185","0.846","0.185"),
            ("0.852","0.182","0.850","0.183","0.846","0.186","0.844","0.187"),
            ("0.859","0.177","0.850","0.183","0.844","0.187","0.844","0.187"),
            ("0.861","0.176","0.848","0.184","0.844","0.186","0.842","0.188"),
        ],
        2: [
            ("0.838","0.190","0.841","0.188","0.840","0.189","0.837","0.190"),
            ("0.839","0.189","0.843","0.187","0.843","0.187","0.840","0.189"),
            ("0.830","0.194","0.829","0.195","0.830","0.195","0.829","0.195"),
            ("0.837","0.191","0.835","0.192","0.835","0.191","0.834","0.192"),
            ("0.834","0.192","0.839","0.189","0.843","0.187","0.839","0.189"),
            ("0.840","0.189","0.841","0.188","0.842","0.187","0.843","0.187"),
            ("0.836","0.191","0.843","0.187","0.844","0.186","0.840","0.189"),
            ("0.841","0.188","0.847","0.184","0.848","0.184","0.846","0.185"),
            ("0.848","0.184","0.849","0.183","0.852","0.181","0.849","0.183"),
        ],
        3: [
            ("0.817","0.202","0.819","0.200","0.818","0.201","0.817","0.202"),
            ("0.818","0.201","0.820","0.200","0.819","0.201","0.818","0.201"),
            ("0.804","0.209","0.807","0.207","0.808","0.207","0.809","0.206"),
            ("0.811","0.205","0.814","0.204","0.813","0.204","0.814","0.203"),
            ("0.801","0.210","0.805","0.208","0.814","0.203","0.817","0.201"),
            ("0.807","0.207","0.818","0.201","0.832","0.193","0.831","0.194"),
            ("0.814","0.203","0.814","0.204","0.819","0.200","0.818","0.201"),
            ("0.810","0.206","0.815","0.203","0.822","0.199","0.826","0.197"),
            ("0.834","0.192","0.831","0.194","0.833","0.193","0.831","0.194"),
        ],
        4: [
            ("0.792","0.215","0.794","0.215","0.793","0.215","0.792","0.215"),
            ("0.793","0.215","0.795","0.214","0.793","0.215","0.793","0.215"),
            ("0.783","0.220","0.784","0.219","0.786","0.219","0.785","0.219"),
            ("0.788","0.218","0.789","0.217","0.789","0.217","0.788","0.217"),
            ("0.714","0.253","0.735","0.243","0.737","0.242","0.747","0.237"),
            ("0.733","0.244","0.747","0.237","0.771","0.226","0.772","0.226"),
            ("0.747","0.238","0.761","0.231","0.770","0.226","0.775","0.224"),
            ("0.773","0.225","0.767","0.228","0.778","0.222","0.778","0.223"),
            ("0.793","0.215","0.796","0.213","0.799","0.212","0.802","0.210"),
        ],
        5: [
            ("0.745","0.239","0.747","0.238","0.749","0.237","0.751","0.236"),
            ("0.743","0.240","0.744","0.239","0.745","0.239","0.748","0.237"),
            ("0.751","0.236","0.752","0.235","0.756","0.233","0.758","0.233"),
            ("0.754","0.235","0.755","0.234","0.758","0.233","0.759","0.232"),
            ("0.651","0.279","0.671","0.271","0.679","0.268","0.674","0.270"),
            ("0.637","0.285","0.684","0.266","0.699","0.259","0.684","0.265"),
            ("0.674","0.270","0.709","0.255","0.723","0.249","0.696","0.261"),
            ("0.680","0.268","0.693","0.262","0.704","0.257","0.702","0.258"),
            ("0.749","0.237","0.751","0.236","0.748","0.237","0.763","0.230"),
        ],
    }
    for h in [1, 2, 3, 4, 5]:
        add_para(doc, f"Horizon +{h} Jam", bold=True, font_size=10)
        rows = []
        for m_idx, model_name in enumerate(_sm_models):
            raw_nse, raw_rmse, mean3_nse, mean3_rmse, med3_nse, med3_rmse, med5_nse, med5_rmse = _sm_data[h][m_idx]
            rows.append([model_name, raw_nse, raw_rmse, mean3_nse, mean3_rmse, med3_nse, med3_rmse, med5_nse, med5_rmse])
        add_table(doc,
            ["Model", "RAW NSE", "RAW RMSE", "MEAN3 NSE", "MEAN3 RMSE", "MED3 NSE", "MED3 RMSE", "MED5 NSE", "MED5 RMSE"],
            rows,
        )

    add_para(doc,
        "Kesimpulan: Smoothing memberikan perbaikan kecil di horizon panjang untuk beberapa model "
        "(MED5 di h5 CatBoost: +0.014) tetapi menurunkan akurasi di h1 secara umum. "
        "Konfigurasi RAW tetap menjadi pilihan paling stabil secara keseluruhan."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 5. DIAGNOSTIK ERROR PER REGIME
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Diagnostik Error: Analisis per Regime Muka Air")
    add_para(doc,
        "Analisis residual dipecah berdasarkan tiga regime muka air untuk mengidentifikasi "
        "kelemahan spesifik model:"
    )
    add_table(doc,
        ["Regime", "Definisi", "Kondisi"],
        [
            ["Normal",   "y < 10 m",       "Operasi harian, base flow"],
            ["Elevated", "10 \u2264 y < 12 m", "Waspada, transisi ke banjir"],
            ["Flood",    "y \u2265 12 m",       "Banjir aktif"],
        ],
    )

    add_heading(doc, "5.1 RMSE per Regime (CatBoost, Test Set)", level=2)
    add_table(doc,
        ["Horizon", "Normal (n=629)", "Elevated (n=61)", "Rasio"],
        [
            ["h1", "0.133 m", "0.409 m", "3.1\u00d7"],
            ["h2", "0.145 m", "0.409 m", "2.8\u00d7"],
            ["h3", "0.156 m", "0.410 m", "2.6\u00d7"],
            ["h4", "0.182 m", "0.425 m", "2.3\u00d7"],
            ["h5", "0.198 m", "0.465 m", "2.4\u00d7"],
        ],
    )
    add_note(doc, "Penting",
        "Error pada regime elevated 2\u20133 kali lipat lebih besar dibanding normal. "
        "Model belum optimal dalam menangkap dinamika transisi dari kondisi normal ke banjir."
    )

    add_heading(doc, "5.2 Bias Sistematik", level=2)
    add_para(doc,
        "Pada regime elevated, terdapat bias negatif (under-prediction) yang konsisten "
        "sebesar \u22120.25 s.d. \u22120.26 m. Model secara konsisten memprediksi lebih rendah "
        "dari aktual. Masalah utama ada pada kemampuan model mengantisipasi kenaikan "
        "ke level elevated/flood."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 6. FLOOD EVENT CROSS-VALIDATION
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Evaluasi Khusus: Flood Event Cross-Validation")

    add_heading(doc, "6.1 Metodologi", level=2)
    add_para(doc,
        "Karena test set temporal (20% terakhir 2023) tidak mengandung flood events besar, "
        "dikembangkan framework evaluasi khusus menggunakan Leave-One-Out Cross-Validation "
        "(LOOCV) pada event banjir:"
    )
    for t in [
        "Deteksi event: Identifikasi episode banjir besar (peak \u2265 13 m, durasi \u2265 10 timestep di atas 12 m).",
        "Buffer exclusion: Hapus window \u00b124 jam di sekitar peak dari data training (mencegah leakage via rolling features 12 jam).",
        "Test window: Evaluasi prediksi pada \u00b16 jam di sekitar peak.",
        "LOOCV: Untuk tiap event, train pada semua data kecuali buffer event tersebut, lalu test pada window event tersebut.",
    ]:
        doc.add_paragraph(t, style="List Number")
    add_para(doc, "Terdeteksi 8 major flood events pada gabungan data 2022\u20132023.")

    add_heading(doc, "6.2 Hasil Flood Event CV (CatBoost)", level=2)
    add_table(doc,
        ["Horizon", "RMSE Mean", "NSE Mean", "Peak Error Mean", "|Peak Error| Mean"],
        [
            ["h1", "0.328 m", "0.915", "\u22120.379 m", "0.379 m"],
            ["h2", "0.425 m", "0.831", "\u22120.401 m", "0.401 m"],
            ["h3", "0.519 m", "0.705", "\u22120.410 m", "0.416 m"],
            ["h4", "0.588 m", "0.624", "\u22120.738 m", "0.738 m"],
            ["h5", "0.625 m", "0.634", "\u22120.755 m", "0.820 m"],
        ],
    )
    add_note(doc, "Temuan Kritis",
        "Semua horizon menunjukkan peak error negatif, artinya model secara konsisten memprediksi "
        "puncak banjir lebih rendah dari aktual. Under-prediction berkisar 0.4 m (h1\u2013h3) "
        "hingga 0.7\u20130.8 m (h4\u2013h5). Untuk sistem EWS, ini berarti peringatan banjir "
        "cenderung terlambat atau kurang urgent."
    )

    add_heading(doc, "6.3 Peak-Weighted RMSE pada Temporal Test Set", level=2)
    add_para(doc,
        "Metrik tambahan yang flood-sensitive, di mana sampel flood diberi bobot 10\u00d7 "
        "dan elevated 3\u00d7:"
    )
    add_table(doc,
        ["Horizon", "RMSE", "wRMSE", "NSE"],
        [
            ["h1", "0.176", "0.227", "0.861"],
            ["h2", "0.184", "0.232", "0.848"],
            ["h3", "0.192", "0.238", "0.834"],
            ["h4", "0.215", "0.257", "0.793"],
            ["h5", "0.237", "0.280", "0.749"],
        ],
    )
    add_para(doc,
        "wRMSE ~30% lebih tinggi dari RMSE biasa, mengkonfirmasi bahwa error terkonsentrasi "
        "pada sampel-sampel dengan muka air tinggi."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 7. PEAK-WEIGHTED TRAINING
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Eksperimen Peak-Weighted Training")

    add_heading(doc, "7.1 Motivasi dan Desain", level=2)
    add_para(doc,
        "Untuk mengatasi systematic peak under-prediction, diuji pendekatan sample weighting "
        "saat training. Dengan memberi bobot lebih besar pada sampel flood, model diharapkan "
        "mempelajari pola peak lebih akurat."
    )
    add_table(doc,
        ["Skema", "Normal (y<10)", "Elevated (10\u201312)", "Flood (y\u226512)", "Karakteristik"],
        [
            ["W0 (baseline)",   "1", "1",           "1",           "Tanpa weighting"],
            ["W1 (moderate)",   "1", "3",           "10",          "Boost moderat"],
            ["W2 (aggressive)", "1", "5",           "20",          "Boost agresif"],
            ["W3 (quadratic)",  "1", "1 + (y\u22129)\u00b2", "1 + (y\u22129)\u00b2", "Continuous boost"],
        ],
    )

    add_heading(doc, "7.2 Hasil pada Flood Event CV", level=2)
    add_table(doc,
        ["Skema", "h1 RMSE", "h1 |Peak Err|", "h3 RMSE", "h3 |Peak Err|", "h5 RMSE", "h5 |Peak Err|", "h5 NSE"],
        [
            ["W0", "0.328", "0.379 m", "0.519", "0.416 m", "0.625", "0.820 m", "0.634"],
            ["W1", "0.329", "0.343 m", "0.509", "0.450 m", "0.616", "0.782 m", "0.648"],
            ["W2", "0.325", "0.338 m", "0.506", "0.471 m", "0.616", "0.734 m", "0.650"],
            ["W3", "0.337", "0.347 m", "0.493", "0.456 m", "0.614", "0.731 m", "0.651"],
        ],
    )

    add_heading(doc, "7.3 Dampak pada Normal Operations", level=2)
    add_table(doc,
        ["Skema", "h1 NSE", "h1 RMSE", "h3 NSE", "h3 RMSE", "h5 NSE", "h5 RMSE", "h5 PBIAS"],
        [
            ["W0", "0.861", "0.176", "0.834", "0.192", "0.749", "0.237", "0.43%"],
            ["W1", "0.840", "0.189", "0.795", "0.214", "0.654", "0.278", "1.00%"],
            ["W2", "0.825", "0.198", "0.772", "0.225", "0.556", "0.315", "1.43%"],
            ["W3", "0.839", "0.189", "0.789", "0.217", "0.599", "0.299", "1.31%"],
        ],
    )
    add_note(doc, "Penting",
        "Peak-weighted training memberikan perbaikan kecil pada peak error (W2 memperbaiki |peak error| "
        "h1 dari 0.379 \u2192 0.338 m), tetapi mengorbankan NSE normal operations secara signifikan. "
        "W2 di h5: NSE turun dari 0.749 \u2192 0.556 (penurunan 0.193). Karena ~90% waktu operasional berada "
        "di regime normal, degradasi ini tidak dapat diterima. Pendekatan sample weighting tidak "
        "direkomendasikan sebagai solusi tunggal."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 8. RINGKASAN MODEL TERBAIK
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Ringkasan Model Terbaik")

    add_heading(doc, "8.1 Konfigurasi Final", level=2)
    add_table(doc,
        ["Parameter", "Nilai"],
        [
            ["Algoritma",  "CatBoost (h1\u2013h4), ElasticNet (h5)"],
            ["Fitur",      "Baseline A (160 fitur + rainfall)"],
            ["Target",     "Nilai absolut muka air (bukan delta)"],
            ["Smoothing",  "Tanpa smoothing (RAW)"],
            ["Weighting",  "Tanpa sample weighting (uniform)"],
            ["Scaling",    "StandardScaler hanya untuk model linear"],
        ],
    )

    add_heading(doc, "8.2 Performa Final", level=2)
    add_table(doc,
        ["Horizon", "Lead Time", "NSE", "RMSE (m)", "MAE (m)", "wRMSE (m)", "Peak Err (m)"],
        [
            ["h1", "+1 jam", "0.861", "0.176", "0.099", "0.227", "\u22120.379"],
            ["h2", "+2 jam", "0.848", "0.184", "0.111", "0.232", "\u22120.401"],
            ["h3", "+3 jam", "0.834", "0.192", "0.120", "0.238", "\u22120.410"],
            ["h4", "+4 jam", "0.793", "0.215", "0.134", "0.257", "\u22120.738"],
            ["h5", "+5 jam", "0.749", "0.237", "0.154", "0.280", "\u22120.755"],
        ],
    )

    add_heading(doc, "8.3 Zona Operasional", level=2)
    add_table(doc,
        ["Horizon", "NSE", "Klasifikasi", "Rekomendasi Penggunaan"],
        [
            ["h1\u2013h2", "0.85\u20130.86", "Reliable", "Keputusan taktis evakuasi dan perencanaan respons"],
            ["h3",         "0.83",            "Good", "Peringatan dini"],
            ["h4",         "0.79",            "Fair", "Estimasi awal, perlu konfirmasi"],
            ["h5",         "0.75",            "Satisfactory", "Estimasi kasar, kombinasikan dengan expert judgement"],
        ],
    )

    add_heading(doc, "8.4 Limitasi yang Teridentifikasi", level=2)
    for t in [
        "Systematic peak under-prediction: Model secara konsisten memprediksi puncak banjir 0.4\u20130.8 m lebih rendah dari aktual.",
        "Degradasi horizon panjang: NSE menurun dari 0.861 (h1) ke 0.749 (h5), dengan penurunan tajam di h4.",
        "Gap regime: Error pada kondisi elevated/flood 2\u20133\u00d7 lebih besar dibanding kondisi normal.",
        "Data curah hujan terbatas: Curah hujan hanya tersedia pada dataset 2023, membatasi pemanfaatan sinyal rainfall.",
    ]:
        doc.add_paragraph(t, style="List Number")

    # ═══════════════════════════════════════════════════════════════════════
    # 9. STRATEGI PENANGANAN MISSING DATA (IMPUTASI)
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Strategi Penanganan Missing Data (Imputasi)")

    add_heading(doc, "9.1 Latar Belakang Masalah", level=2)
    add_para(doc,
        "Dalam operasi nyata, sensor muka air dapat mengalami gangguan (kerusakan, gangguan komunikasi, "
        "pemeliharaan). Saat satu atau beberapa stasiun tidak mengirim data, seluruh fitur yang diturunkan "
        "dari stasiun tersebut (t0, lag, rolling mean/std, diff) menjadi kosong. Model saat ini tidak "
        "dapat menghasilkan prediksi jika ada fitur yang bernilai NaN."
    )
    add_para(doc,
        "Dengan 13 stasiun yang masing-masing menghasilkan ~12 fitur, hilangnya 1 stasiun berarti "
        "~12 fitur (7.5% dari total) menjadi tidak tersedia."
    )

    add_heading(doc, "9.2 Klasifikasi Skenario Missing Data", level=2)
    add_table(doc,
        ["Skenario", "Deskripsi", "Frekuensi", "Dampak"],
        [
            ["S1: Downtime singkat",    "1 stasiun mati < 2 jam",          "Sering",        "Rendah"],
            ["S2: Downtime sedang",     "1 stasiun mati 2\u201312 jam",    "Kadang",        "Sedang"],
            ["S3: Multi-stasiun",       "2\u20133 stasiun mati bersamaan", "Jarang",        "Tinggi"],
            ["S4: Stasiun kritikal",    "Dhompo sendiri mati",             "Sangat jarang", "Sangat tinggi"],
        ],
    )

    add_heading(doc, "9.3 Metode Imputasi", level=2)

    add_heading(doc, "A. Forward-Fill Temporal (untuk S1: downtime singkat)", level=3)
    add_para(doc,
        "Metode paling sederhana: gunakan nilai terakhir yang valid dari stasiun yang bermasalah. "
        "Muka air berubah relatif lambat (~0.02\u20130.05 m per 30 menit pada kondisi normal), sehingga "
        "forward-fill masih representatif untuk gap pendek. Disarankan maksimal 4 timestep (2 jam)."
    )

    add_heading(doc, "B. Interpolasi Spasial (untuk S2: downtime sedang)", level=3)
    add_para(doc,
        "Memanfaatkan korelasi spasial antar stasiun yang berdekatan secara elevasi. "
        "Nilai stasiun yang hilang diestimasi dari stasiun tetangga menggunakan "
        "inverse distance weighting berdasarkan elevasi."
    )
    add_para(doc,
        "Contoh: Jika Bd. Bakalan (136 m) mati, estimasi dari Bd. Lecari (167 m) dengan jarak "
        "elevasi 31 m dan AWLR Kademungan (128 m) dengan jarak elevasi 8 m. Stasiun dengan "
        "jarak elevasi lebih kecil mendapat bobot lebih besar karena korelasi hidrologis yang lebih tinggi "
        "(dikonfirmasi dari analisis EDA)."
    )

    add_heading(doc, "C. Kombinasi Forward-Fill + Spatial (untuk S3: multi-stasiun)", level=3)
    add_para(doc, "Jika beberapa stasiun hilang bersamaan:")
    for t in [
        "Prioritas forward-fill untuk stasiun dengan gap < 2 jam.",
        "Interpolasi spasial untuk stasiun dengan gap lebih panjang, menggunakan stasiun yang masih aktif sebagai referensi.",
        "Jika stasiun tetangga terdekat juga mati, perluas pencarian ke tetangga berikutnya berdasarkan ranking elevasi.",
    ]:
        doc.add_paragraph(t, style="List Number")

    add_heading(doc, "D. Penanganan Khusus Stasiun Dhompo (untuk S4)", level=3)
    add_para(doc, "Jika Dhompo sendiri (target station) mati:")
    for t in [
        "Forward-fill maksimal 1 jam (2 timestep) karena Dhompo_t0 adalah fitur paling dominan.",
        "Jika gap > 1 jam, estimasi dari stasiun hilir terdekat (Klosod, elevasi 22 m) dengan koreksi offset berdasarkan selisih rata-rata historis.",
        'Flag confidence = "degraded" pada output prediksi agar pengguna mengetahui bahwa prediksi menggunakan data estimasi.',
    ]:
        doc.add_paragraph(t, style="List Number")

    add_heading(doc, "9.4 Penanganan Fitur Turunan", level=2)
    add_para(doc,
        "Imputasi pada nilai raw stasiun saja tidak cukup \u2014 fitur turunan juga perlu diperbarui:"
    )
    add_table(doc,
        ["Fitur", "Penanganan"],
        [
            ["t0, lag1\u20133",    "Diisi dari nilai imputasi"],
            ["Rolling mean", "Dihitung ulang dengan data valid + imputasi (min_periods lebih rendah dari window)"],
            ["Rolling std",  "Sama dengan rolling mean; jika window terlalu pendek, gunakan nilai std dari timestep sebelumnya"],
            ["diff1, diff2", "Dihitung dari nilai imputasi; jika timestep sebelumnya juga imputasi, set ke 0 (asumsi tidak ada perubahan)"],
        ],
    )

    add_heading(doc, "9.5 Confidence Scoring", level=2)
    add_para(doc,
        "Setiap prediksi disertai skor confidence yang merefleksikan kualitas data input:"
    )
    add_table(doc,
        ["Kondisi", "Confidence", "Keterangan"],
        [
            ["Semua stasiun aktif",                                   "High",     "Prediksi penuh"],
            ["1\u20132 stasiun non-kritikal imputasi, gap < 2 jam",   "Medium",   "Prediksi masih reliable"],
            ["\u22653 stasiun imputasi, atau 1 stasiun kritikal",     "Low",      "Prediksi perlu verifikasi manual"],
            ["Dhompo imputasi > 1 jam, atau >50% stasiun mati",       "Degraded", "Tidak direkomendasikan untuk keputusan kritis"],
        ],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 10. REKOMENDASI STRATEGI SELANJUTNYA
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Rekomendasi Strategi Selanjutnya")

    add_para(doc,
        "Berdasarkan temuan dari seluruh eksperimen, berikut strategi pengembangan yang direkomendasikan:"
    )

    add_heading(doc, "10.1 Regime-Switching Model", level=2)
    add_para(doc,
        "Membangun dua model terpisah untuk regime normal dan regime flood, dengan classifier atau "
        "threshold untuk switching. Masing-masing model dapat memiliki feature importance dan "
        "hyperparameter yang dioptimasi untuk regime-nya."
    )

    add_heading(doc, "10.2 Asymmetric Loss / Quantile Regression", level=2)
    add_para(doc,
        "Mengubah fungsi loss agar lebih berat menghukum under-prediction daripada over-prediction. "
        "CatBoost mendukung Quantile:alpha=0.6 yang akan menggeser prediksi sedikit ke atas saat banjir, "
        "tanpa perlu mengubah data atau bobot sampel."
    )

    add_heading(doc, "10.3 Post-Processing Peak Correction", level=2)
    add_para(doc,
        "Menambahkan lapisan koreksi bias khusus regime flood: y_corrected = y_pred + f(y_pred) "
        "di mana f() hanya aktif saat y_pred melampaui threshold. Ini tidak memerlukan retraining model."
    )

    add_heading(doc, "10.4 Sequence Model untuk Horizon Panjang", level=2)
    add_para(doc,
        "LSTM atau Transformer untuk horizon h4\u2013h5 yang memerlukan pemahaman temporal dynamics "
        "lebih dalam. Dapat dikombinasikan dengan model CatBoost untuk horizon pendek dalam "
        "arsitektur ensemble."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # REFERENSI OUTPUT
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "Referensi Output")
    add_table(doc,
        ["File", "Isi"],
        [
            ["xls_11_model_comparison_final.xlsx",       "Perbandingan lengkap semua model"],
            ["xls_13_combined_experiment_comparison.xlsx","Eksperimen A vs B vs C"],
            ["experiment_progressive_features.xlsx",     "Progressive features A\u2013B4"],
            ["experiment_delta_vs_abs.xlsx",             "Delta vs absolute target"],
            ["experiment_target_smoothing.xlsx",         "Target smoothing"],
            ["flood_event_cv.xlsx",                      "Flood event LOOCV detail"],
            ["peak_weighted_rmse.xlsx",                  "Peak-weighted RMSE"],
            ["peak_weighted_flood_cv.xlsx",              "Peak-weighted training \u00d7 Flood CV"],
            ["peak_weighted_normops.xlsx",               "Peak-weighted training \u00d7 NormOps"],
            ["diagnostic_regime_errors.xlsx",            "Error per regime"],
        ],
    )
    add_para(doc,
        "Semua file tabel tersimpan di direktori reports/tables/.",
        italic=True, font_size=10,
    )

    return doc


# ── Entry point ──────────────────────────────────────────────────────────────

def main():
    out_path = PROJECT_ROOT / "reports" / "laporan_training_model_dhompo.docx"
    doc = build_document()
    doc.save(str(out_path))
    print(f"Laporan berhasil disimpan: {out_path}")


if __name__ == "__main__":
    main()
